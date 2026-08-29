"""Business logic for letters, workflow status updates, documents, and messaging."""

from __future__ import annotations

import base64
import hashlib
import html
import os
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import parse_qs, urlparse
from zoneinfo import ZoneInfo

import qrcode
import requests as _requests
from flask import current_app, has_request_context, jsonify, request, session, url_for
from sqlalchemy import func

from . import settings
from .auth import normalize_email
from .database import ensure_dirs
from .extensions import db
from .models import Letter

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

try:
    from barcode import Code128
    from barcode.writer import ImageWriter

    HAVE_BARCODE = True
except Exception:
    HAVE_BARCODE = False

try:
    from PIL import Image

    HAVE_PIL = True
except Exception:
    HAVE_PIL = False


IST_ZONE = ZoneInfo("Asia/Kolkata")


def utc_now() -> datetime:
    """Return a timezone-aware datetime for audit fields."""
    return datetime.now(timezone.utc)


def to_ist(value: datetime | None) -> datetime | None:
    """Convert stored UTC datetimes into Asia/Kolkata timezone."""
    if value is None:
        return None
    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone.utc)
    return value.astimezone(IST_ZONE)


def serialize_datetime(value: datetime | None) -> str | None:
    """Convert datetimes into stable IST ISO strings for templates and JSON."""
    if value is None:
        return None
    value = to_ist(value)
    return value.replace(microsecond=0).isoformat()


def format_letter_date(value: datetime | None) -> str:
    """Format stored timestamps for the formal letter header."""
    if value is None:
        return datetime.now(IST_ZONE).strftime("%d %B %Y")
    return to_ist(value).strftime("%d %B %Y")


def sentence_case(text: str) -> str:
    """Convert short user-entered subjects into a cleaner display form."""
    value = re.sub(r"\s+", " ", text or "").strip(" .")
    if not value:
        return ""
    return value[0].upper() + value[1:]


def normalize_letter_description(text: str) -> str:
    """Keep the student's wording, but trim noise so the letter stays readable and on one page."""
    paragraphs = []
    for raw_line in (text or "").splitlines():
        cleaned = re.sub(r"\s+", " ", raw_line).strip()
        if cleaned:
            paragraphs.append(cleaned)
    return "\n".join(paragraphs).strip()


def resolve_signature_file(signature_data_or_name: str | None, identifier: str) -> str:
    """Ensure the signature image file exists on disk, reconstructing it from Base64 if needed."""
    if not signature_data_or_name:
        return ""

    ensure_dirs()
    if signature_data_or_name.startswith("data:image"):
        try:
            _, encoded = signature_data_or_name.split(",", 1)
            raw_bytes = base64.b64decode(encoded)
            file_name = f"cached_{identifier}.png"
            target_path = os.path.join(settings.SIGNATURE_DIR, file_name)
            with open(target_path, "wb") as handle:
                handle.write(raw_bytes)
            return target_path
        except Exception as exc:
            if current_app:
                current_app.logger.warning("Failed to decode signature: %s", exc)
            return ""

    disk_path = os.path.join(settings.SIGNATURE_DIR, signature_data_or_name)
    return disk_path if os.path.exists(disk_path) else ""


def build_formal_letter_content(letter: Letter) -> dict[str, object]:
    """Build the fixed-format formal letter content with reliable signature resolution."""
    subject = sentence_case(letter.generated_subject or letter.subject) or "Request"
    body_text = normalize_letter_description(letter.generated_body or letter.description)
    sig_path = resolve_signature_file(letter.signature_file_name, letter.app_id)

    return {
        "heading_line": settings.LETTER_HEADING,
        "from_lines": [
            letter.name,
            letter.email,
            letter.phone,
        ],
        "date_line": format_letter_date(letter.created_at),
        "to_lines": [
            "The Head of the Department",
            settings.LETTER_HEADING,
            "MIT Campus",
            "Anna University",
            settings.CITY_TITLE,
        ],
        "subject_line": subject if subject.lower().endswith("-reg.") else f"{subject} - Reg.",
        "body_text": body_text,
        "hod_verification": letter.status == settings.STATUS_APPROVED,
        "signature_path": sig_path,
    }


def serialize_letter(letter: Letter) -> dict[str, str | None]:
    """Convert a letter model into a JSON-safe structure."""
    return {
        "app_id": letter.app_id,
        "name": letter.name,
        "email": letter.email,
        "phone": letter.phone,
        "request_type": letter.request_type,
        "subject": letter.subject,
        "description": letter.description,
        "original_description": letter.original_description,
        "generated_subject": letter.generated_subject,
        "generated_body": letter.generated_body,
        "status": letter.status,
        "created_at": serialize_datetime(letter.created_at),
        "submitted_at": serialize_datetime(letter.submitted_at),
        "approved_at": serialize_datetime(letter.approved_at),
    }


def is_allowed_institute_email(email: str) -> bool:
    """Restrict public student signup to configured institute domains."""
    email = normalize_email(email)
    if "@" not in email:
        return False
    if not settings.ALLOWED_EMAIL_DOMAINS:
        return True
    return email.rsplit("@", 1)[1] in settings.ALLOWED_EMAIL_DOMAINS


def generate_app_id() -> str:
    """Generate a short unique letter identifier."""
    for _ in range(10):
        app_id = uuid.uuid4().hex[:8]
        if not db.session.get(Letter, app_id):
            return app_id
    raise RuntimeError("Unable to generate a unique application ID")


def build_submit_url(app_id: str) -> str:
    """Build the QR payload URL for status tracking."""
    path = f"/submit?id={app_id}"
    if settings.APP_BASE_URL:
        return f"{settings.APP_BASE_URL}{path}"
    if has_request_context():
        return url_for("submit", id=app_id, _external=True)
    return path


def extract_app_id(text: str | None) -> str | None:
    """Extract the tracked letter ID from URLs, raw codes, or fallback payloads."""
    if not text:
        return None

    value = text.strip()
    if not value:
        return None

    try:
        parsed = urlparse(value)
        params = parse_qs(parsed.query)
        if params.get("id"):
            return params["id"][0].strip()
    except Exception:
        pass

    pipe_parts = [part.strip() for part in value.split("|") if part.strip()]
    if pipe_parts:
        last = pipe_parts[-1]
        if re.fullmatch(r"[A-Za-z0-9_-]{6,32}", last):
            return last

    if re.fullmatch(r"[A-Za-z0-9_-]{6,32}", value):
        return value

    match = re.search(r"[?&]id=([A-Za-z0-9_-]{6,32})", value)
    if match:
        return match.group(1)

    return None


def letter_belongs_to_current_user(letter: Letter) -> bool:
    """Restrict student access to only their own letters."""
    if session.get("role") in {"staff", "admin"}:
        return True
    return session.get("role") == "student" and normalize_email(letter.email) == normalize_email(session.get("email", ""))


def get_letter(app_id: str) -> Letter | None:
    """Fetch a letter by its tracked ID."""
    return db.session.get(Letter, app_id)


def update_letter_status(app_id: str, new_status: str) -> Letter | None:
    """Advance the workflow and stamp important lifecycle dates."""
    if new_status not in settings.VALID_STATUSES:
        raise ValueError("Invalid status")

    letter = db.session.get(Letter, app_id)
    if not letter:
        return None

    current_status = letter.status
    if current_status != new_status and new_status not in settings.STATUS_FLOW.get(current_status, set()):
        raise ValueError(f"Cannot move letter from {current_status} to {new_status}")

    if current_status != new_status:
        timestamp = utc_now()
        letter.status = new_status
        if new_status == settings.STATUS_SUBMITTED and not letter.submitted_at:
            letter.submitted_at = timestamp
        if new_status == settings.STATUS_APPROVED and not letter.approved_at:
            letter.approved_at = timestamp
        db.session.commit()
        notification_status = notify_letter_status_change(letter, current_status, new_status, timestamp)
        letter.notification_status = notification_status
        letter.email_status = notification_status["email"]
        letter.sms_status = notification_status["sms"]

    return letter


def build_artifact_name(app_id: str, extension: str, existing_name: str | None = None) -> str:
    """Create stable unique artifact names for generated files."""
    if existing_name:
        return existing_name
    return f"{app_id}_{uuid.uuid4().hex[:10]}.{extension}"


def create_qr_assets(letter: Letter) -> tuple[str, str | None]:
    """Generate QR and optional Code128 barcode images for a letter."""
    ensure_dirs()
    letter.qr_file_name = build_artifact_name(letter.app_id, "png", letter.qr_file_name)
    qr_path = os.path.join(settings.QR_DIR, letter.qr_file_name)
    payload = build_submit_url(letter.app_id)
    img = qrcode.make(payload)
    img.save(qr_path)

    barcode_path = None
    if HAVE_BARCODE:
        try:
            barcode_filebase = os.path.join(settings.BARCODE_DIR, letter.app_id)
            Code128(payload, writer=ImageWriter()).save(barcode_filebase)
            barcode_path = f"{barcode_filebase}.png"
        except Exception as exc:
            if current_app:
                current_app.logger.warning("Failed to generate barcode for %s: %s", letter.app_id, exc)

    db.session.commit()
    return qr_path, barcode_path


def generate_letter_file(letter: Letter) -> str:
    """Create a compact downloadable letter that works well on both mobile and desktop downloads."""
    ensure_dirs()
    qr_path, _ = create_qr_assets(letter)
    letter_content = build_formal_letter_content(letter)
    extension = "docx" if HAVE_DOCX else "txt"
    letter.generated_file_name = build_artifact_name(letter.app_id, extension, letter.generated_file_name)
    output_path = os.path.join(settings.GEN_DIR, letter.generated_file_name)

    if HAVE_DOCX:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(10.5)

        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run(letter_content["heading_line"])
        heading_run.bold = True
        heading_run.font.size = Pt(13)

        doc.add_paragraph()

        top_table = doc.add_table(rows=1, cols=2)
        top_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        top_table.autofit = False
        left_cell = top_table.cell(0, 0)
        right_cell = top_table.cell(0, 1)

        from_paragraph = left_cell.paragraphs[0]
        from_paragraph.add_run("From:\n").bold = True
        from_paragraph.add_run("\n".join(letter_content["from_lines"]))

        qr_paragraph = right_cell.paragraphs[0]
        qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        qr_paragraph.add_run().add_picture(qr_path, width=Inches(0.9))

        doc.add_paragraph()
        doc.add_paragraph(letter_content["date_line"])
        doc.add_paragraph()

        to_paragraph = doc.add_paragraph()
        to_paragraph.add_run("To,\n").bold = True
        to_paragraph.add_run("\n".join(letter_content["to_lines"]))

        doc.add_paragraph()
        doc.add_paragraph("Respected Sir/Mam,")

        subject_paragraph = doc.add_paragraph()
        subject_paragraph.add_run("Sub").bold = True
        subject_paragraph.add_run(" : ")
        subject_paragraph.add_run(letter_content["subject_line"])

        for paragraph_text in (letter_content["body_text"] or "").splitlines():
            body_paragraph = doc.add_paragraph(paragraph_text)
            body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph()
        doc.add_paragraph("Yours sincerely,")
        if letter_content["signature_path"] and os.path.exists(letter_content["signature_path"]):
            signature_paragraph = doc.add_paragraph()
            signature_paragraph.paragraph_format.space_after = Pt(0)
            signature_paragraph.add_run().add_picture(letter_content["signature_path"], width=Inches(1.35))
        else:
            doc.add_paragraph("")
        doc.add_paragraph(letter.name)

        if letter_content["hod_verification"]:
            verification = doc.add_paragraph()
            verification.alignment = WD_ALIGN_PARAGRAPH.CENTER
            check = verification.add_run("\u2713 ")
            check.bold = True
            check.font.color.rgb = RGBColor(0x16, 0x8A, 0x45)
            verified_text = verification.add_run("Digitally verified by the HoD")
            verified_text.bold = True
            verified_text.font.color.rgb = RGBColor(0x16, 0x8A, 0x45)

        doc.save(output_path)
    else:
        with open(output_path, "w", encoding="utf-8") as handle:
            handle.write(f"{letter_content['heading_line'].center(70)}\n\n")
            handle.write("From:\n")
            for line in letter_content["from_lines"]:
                handle.write(f"{line}\n")
            handle.write(f"\n{letter_content['date_line']}\n\nTo,\n")
            for line in letter_content["to_lines"]:
                handle.write(f"{line}\n")
            handle.write(f"\nRespected Sir/Mam,\nSub : {letter_content['subject_line']}\n\n")
            handle.write(f"{letter_content['body_text']}\n\nYours sincerely,\n")
            handle.write(f"{letter.name}\n")
            if letter_content["hod_verification"]:
                handle.write("\n\u2713 Digitally verified by the HoD\n")

    db.session.commit()
    return output_path


def save_signature_image(file_storage, identifier: str) -> str:
    """Read image bytes, validate format, and return Base64 string for database persistence."""
    if not HAVE_PIL:
        raise ValueError("Image support is unavailable on the server.")

    ensure_dirs()
    file_storage.stream.seek(0)
    image = Image.open(file_storage.stream)
    image.verify()
    file_storage.stream.seek(0)

    image_bytes = file_storage.read()
    b64_string = f"data:image/png;base64,{base64.b64encode(image_bytes).decode('utf-8')}"

    file_name = f"{identifier}_{uuid.uuid4().hex[:8]}.png"
    local_path = os.path.join(settings.SIGNATURE_DIR, file_name)
    with open(local_path, "wb") as handle:
        handle.write(image_bytes)

    return b64_string


def ensure_letter_file(letter: Letter) -> str:
    """Return a valid generated letter file path, recreating it safely when storage is ephemeral."""
    ensure_dirs()
    if letter.status == settings.STATUS_APPROVED:
        return generate_letter_file(letter)
    if letter.generated_file_name:
        current_path = os.path.join(settings.GEN_DIR, letter.generated_file_name)
        if os.path.exists(current_path):
            return current_path
    return generate_letter_file(letter)


def hash_reset_token(token: str) -> str:
    """Hash a reset token before storing it or comparing it to database values."""
    value = (token or "").strip()
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def build_password_reset_link(raw_token: str) -> str:
    """Build an absolute reset URL from config or active request context."""
    base_url = (settings.APP_BASE_URL or os.environ.get("APP_BASE_URL", "")).strip().rstrip("/")
    token_value = (raw_token or "").strip()

    if base_url:
        return f"{base_url}/reset-password/{token_value}"

    if has_request_context():
        return url_for("reset_password", token=token_value, _external=True)

    return f"/reset-password/{token_value}"


def email_is_configured() -> bool:
    """Return whether Mailjet HTTP delivery is configured."""
    return bool(
        os.environ.get("MAILJET_API_KEY")
        and os.environ.get("MAILJET_SECRET_KEY")
        and os.environ.get("MAILJET_SENDER_EMAIL")
    )


def sms_is_configured() -> bool:
    """Return whether optional Twilio SMS delivery is enabled and configured."""
    return bool(
        settings.SMS_ENABLED
        and settings.TWILIO_ACCOUNT_SID
        and settings.TWILIO_AUTH_TOKEN
        and settings.TWILIO_FROM_NUMBER
    )


def normalize_sms_phone(phone: str | None) -> str:
    """Convert a stored student phone number to E.164 form for Twilio."""
    value = re.sub(r"[\s()\-]", "", phone or "")
    if value.startswith("00"):
        value = "+" + value[2:]
    elif not value.startswith("+"):
        if value.isdigit() and len(value) == 10:
            value = f"{settings.SMS_DEFAULT_COUNTRY_CODE}{value}"
        else:
            return ""
    if not re.fullmatch(r"\+[1-9]\d{7,14}", value):
        return ""
    return value


def send_sms(phone: str | None, message: str, ref: str | None = None) -> bool:
    """Send an optional status SMS via Twilio without storing provider credentials in code."""
    if not sms_is_configured():
        return False

    recipient = normalize_sms_phone(phone)
    if not recipient:
        return False

    url = f"https://api.twilio.com/2010-04-01/Accounts/{settings.TWILIO_ACCOUNT_SID}/Messages.json"
    try:
        response = _requests.post(
            url,
            data={"To": recipient, "From": settings.TWILIO_FROM_NUMBER, "Body": message},
            auth=(settings.TWILIO_ACCOUNT_SID, settings.TWILIO_AUTH_TOKEN),
            timeout=settings.SMS_TIMEOUT,
        )
        return response.status_code in {200, 201}
    except Exception:
        return False


def notification_summary(notification_status: dict[str, str] | None) -> str:
    """Create an honest, concise delivery message for HTML and device clients."""
    status = notification_status or {}
    channel_labels = {"email": "Email", "sms": "SMS"}
    outcomes = []
    for channel in ("email", "sms"):
        outcome = status.get(channel)
        if outcome == "sent":
            outcomes.append(f"{channel_labels[channel]} sent")
        elif outcome == "failed":
            outcomes.append(f"{channel_labels[channel]} failed")
    if outcomes:
        return "; ".join(outcomes) + "."
    return "No notification channel is configured."


def plain_text_to_html(content: str) -> str:
    """Safely wrap plain notification text for Mailjet's HTMLPart field."""
    paragraphs = [
        html.escape(paragraph).replace("\n", "<br>\n")
        for paragraph in (content or "").strip().split("\n\n")
        if paragraph.strip()
    ]
    body = "\n".join(f"<p>{paragraph}</p>" for paragraph in paragraphs)
    return f"<!doctype html><html><body>{body}</body></html>"


def notify_letter_status_change(
    letter: Letter,
    previous_status: str,
    new_status: str,
    updated_at: datetime | None = None,
) -> dict[str, str]:
    """Notify a student when a letter status changes without repeating unchanged values."""
    result = {
        "email": "not_configured" if not email_is_configured() else "failed",
        "sms": "not_configured" if not sms_is_configured() else "failed",
    }
    if not letter or not letter.email or previous_status == new_status:
        return result

    updated_at = updated_at or utc_now()
    portal_url = f"{settings.APP_BASE_URL}/submit?id={letter.app_id}" if settings.APP_BASE_URL else f"/submit?id={letter.app_id}"
    remarks = (letter.description or "").strip() or "No remarks were added."
    body = (
        f"Dear {letter.name or 'Student'},\n\n"
        f"Your letter status has been updated.\n\n"
        f"Letter ID: {letter.app_id}\n"
        f"Subject: {letter.subject}\n"
        f"Previous Status: {previous_status}\n"
        f"New Status: {new_status}\n"
        f"Updated Time (IST): {serialize_datetime(updated_at)}\n"
        f"Remarks / Comments: {remarks}\n"
        f"Track your letter: {portal_url}\n\n"
        "If you did not request this change, you can ignore this email."
    )
    if new_status in {settings.STATUS_SUBMITTED, settings.STATUS_APPROVED} and email_is_configured():
        result["email"] = "sent" if send_mailjet_email(
            letter.email,
            f"Letter Status Updated - {letter.app_id}",
            plain_text_to_html(body),
            ref=letter.app_id,
        ) else "failed"

    if new_status in {settings.STATUS_SUBMITTED, settings.STATUS_APPROVED} and sms_is_configured():
        sms_body = f"MIT Letterbox: {letter.app_id} is now {new_status}. Track: {portal_url}"
        result["sms"] = "sent" if send_sms(letter.phone, sms_body, ref=letter.app_id) else "failed"
    return result


def notify_letter_created(letter: Letter) -> dict[str, str]:
    """Send the initial creation notification through each configured channel."""
    result = {
        "email": "not_configured" if not email_is_configured() else "failed",
        "sms": "not_configured" if not sms_is_configured() else "failed",
    }
    if not letter:
        return result

    portal_url = f"{settings.APP_BASE_URL}/submit?id={letter.app_id}" if settings.APP_BASE_URL else f"/submit?id={letter.app_id}"
    if email_is_configured() and letter.email:
        body = (
            f"Dear {letter.name or 'Student'},\n\n"
            f"Your letter request {letter.app_id} has been created. Download it, print it, and submit it to the institute letterbox.\n\n"
            f"Track your letter: {portal_url}"
        )
        result["email"] = "sent" if send_mailjet_email(
            letter.email,
            f"Letter Created - {letter.app_id}",
            plain_text_to_html(body),
            ref=letter.app_id,
        ) else "failed"
    if sms_is_configured():
        result["sms"] = "sent" if send_sms(
            letter.phone,
            f"MIT Letterbox: Your letter {letter.app_id} was created. Track: {portal_url}",
            ref=letter.app_id,
        ) else "failed"
    return result


def send_mailjet_email(to_email: str, subject: str, html_content: str, ref: str | None = None) -> bool:
    """Send an email through Mailjet's v3.1 HTTP API and keep a local audit copy."""
    recipient = normalize_email(to_email)
    if not recipient or "@" not in recipient or not email_is_configured():
        return False

    plain_content = re.sub(r"<[^>]+>", " ", html_content or "")

    try:
        ensure_dirs()
        stamp = datetime.now().strftime("%Y%m%dT%H%M%S")
        ref_part = f"{ref}_" if ref else ""
        file_name = f"{ref_part}{stamp}_{uuid.uuid4().hex[:8]}.eml"
        file_path = os.path.join(settings.SENT_DIR, file_name)
        with open(file_path, "w", encoding="utf-8") as handle:
            handle.write(f"From: {os.environ.get('MAILJET_SENDER_EMAIL')}\nTo: {recipient}\nSubject: {subject}\n\n{plain_content}")
    except Exception:
        pass

    try:
        response = _requests.post(
            settings.MAILJET_API_URL,
            auth=(os.environ.get("MAILJET_API_KEY"), os.environ.get("MAILJET_SECRET_KEY")),
            json={
                "Messages": [
                    {
                        "From": {
                            "Email": os.environ.get("MAILJET_SENDER_EMAIL"),
                            "Name": "Letter Tracking System",
                        },
                        "To": [{"Email": recipient}],
                        "Subject": subject,
                        "HTMLPart": html_content,
                    }
                ]
            },
            timeout=settings.MAILJET_TIMEOUT,
        )
        return 200 <= response.status_code < 300
    except Exception:
        return False


def jsonify_error(message: str, status_code: int = 400, **payload):
    """Return a consistent JSON error body for AJAX and ESP clients."""
    body = {"status": "error", "message": message}
    body.update(payload)
    return jsonify(body), status_code


def verify_recaptcha(form_token: str | None) -> tuple[bool, str | None]:
    """Verify a Google reCAPTCHA response when keys are configured."""
    if not settings.RECAPTCHA_ENABLED:
        return True, None

    token = (form_token or "").strip()
    if not token:
        return False, "Please complete the reCAPTCHA challenge."

    remote_ip = request.headers.get("X-Forwarded-For", request.remote_addr or "")
    payload = {
        "secret": settings.RECAPTCHA_SECRET_KEY,
        "response": token,
        "remoteip": remote_ip.split(",")[0].strip(),
    }
    try:
        response = _requests.post("https://www.google.com/recaptcha/api/siteverify", data=payload, timeout=8)
        data = response.json()
        if data.get("success"):
            return True, None
        return False, "reCAPTCHA verification failed. Please try again."
    except Exception:
        return False, "Unable to verify reCAPTCHA right now. Please try again."


def esp_token_valid(data) -> bool:
    """Validate the optional ESP device token for hardware endpoints."""
    expected = os.environ.get("ESP_TOKEN", "").strip()
    if not expected:
        return False

    token = request.headers.get("X-ESP-Token", "").strip()
    if not token and hasattr(data, "get"):
        token = str(data.get("token", "")).strip()
    return token == expected


def parse_esp_payload() -> dict[str, str]:
    """Read JSON, form-encoded, or query-string POST data from either ESP board."""
    payload = request.get_json(silent=True)
    if not isinstance(payload, dict):
        payload = request.form.to_dict(flat=True)
    if not payload:
        payload = request.args.to_dict(flat=True)

    normalized = {
        str(key): str(value).strip()
        for key, value in payload.items()
        if value is not None
    }
    scanned_value = (
        normalized.get("id")
        or normalized.get("app_id")
        or normalized.get("letter_id")
        or normalized.get("code")
        or normalized.get("barcode")
        or normalized.get("qr")
    )
    app_id = extract_app_id(scanned_value)
    if app_id:
        normalized["app_id"] = app_id
    return normalized


def infer_esp_action(data) -> str | None:
    """Infer submit/approve actions from explicit action or ESP device identity."""
    action = ""
    if hasattr(data, "get"):
        action = str(data.get("action", "")).strip().lower()
        if action in {"submit", "approve"}:
            return action

        device_id = str(data.get("device_id", "") or data.get("device", "") or data.get("node_id", "") or "").strip()
        if device_id:
            inbox = settings.ESP_INBOX_DEVICE_ID.lower()
            outbox = settings.ESP_OUTBOX_DEVICE_ID.lower()
            if inbox and device_id.lower() == inbox:
                return "submit"
            if outbox and device_id.lower() == outbox:
                return "approve"
    return None